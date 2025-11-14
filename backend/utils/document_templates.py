"""
Document Template Engine for generating tender submission documents
Supports DOCX and PDF generation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from typing import Dict, Any, List
import os
from datetime import datetime

class DocumentTemplateEngine:
    """Generate professional tender documents"""
    
    def __init__(self, output_dir: str = "/tmp"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_cover_letter(self, data: Dict[str, Any]) -> str:
        """
        Generate cover letter DOCX
        
        Args:
            data: Dictionary with tender and company details
            
        Returns:
            Path to generated document
        """
        doc = Document()
        
        # Company letterhead (placeholder)
        header = doc.add_paragraph()
        header.add_run(data.get('company_name', 'Company Name')).bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph()
        
        # To address
        doc.add_paragraph(f"To,")
        doc.add_paragraph(f"{data.get('tender_organization', 'The Tender Inviting Authority')}")
        doc.add_paragraph(f"{data.get('tender_department', '')}")
        doc.add_paragraph()
        
        # Subject
        subject = doc.add_paragraph()
        subject.add_run(f"Subject: Submission of Bid for Tender No. {data.get('tender_number', 'N/A')}").bold = True
        doc.add_paragraph()
        
        # Body
        doc.add_paragraph("Dear Sir/Madam,")
        doc.add_paragraph()
        
        body_text = f"""
        We are pleased to submit our bid for the tender "{data.get('tender_title', '')}" 
        (Tender No: {data.get('tender_number', 'N/A')}).
        
        We have carefully reviewed all the terms and conditions specified in the tender document 
        and confirm our compliance with all the requirements. We have submitted all mandatory 
        documents as required.
        
        Our company, {data.get('company_name', '')}, has extensive experience in similar projects 
        and we are confident of our ability to deliver as per the specifications.
        
        Please find enclosed the following documents:
        1. Technical Bid
        2. Financial Bid (in sealed envelope)
        3. Compliance Statement
        4. Supporting Documents and Annexures
        
        We look forward to your favorable consideration of our bid.
        """
        
        doc.add_paragraph(body_text)
        doc.add_paragraph()
        
        # Closing
        doc.add_paragraph("Thanking you,")
        doc.add_paragraph("Yours faithfully,")
        doc.add_paragraph()
        doc.add_paragraph()
        
        signature = doc.add_paragraph()
        signature.add_run("[Signature]").italic = True
        doc.add_paragraph(f"{data.get('authorized_person', 'Authorized Signatory')}")
        doc.add_paragraph(f"{data.get('designation', 'Director')}")
        doc.add_paragraph(f"{data.get('company_name', '')}")
        
        # Save
        filename = f"cover_letter_{data.get('tender_number', 'tender').replace('/', '_')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_boq_document(self, data: Dict[str, Any]) -> str:
        """
        Generate BOQ document with pricing
        
        Args:
            data: BOQ data with line items
            
        Returns:
            Path to generated document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('Bill of Quantities', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Tender No: {data.get('tender_number', 'N/A')}")
        doc.add_paragraph(f"BOQ No: {data.get('boq_number', 'N/A')}")
        doc.add_paragraph()
        
        # Table
        line_items = data.get('line_items', [])
        if line_items:
            # Create table with headers
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'S.No'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Specification'
            hdr_cells[3].text = 'Qty'
            hdr_cells[4].text = 'Unit'
            hdr_cells[5].text = 'Rate (₹)'
            hdr_cells[6].text = 'Amount (₹)'
            
            # Add data rows
            for i, item in enumerate(line_items, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = item.get('description', '')
                row_cells[2].text = item.get('specification', '')
                row_cells[3].text = str(item.get('quantity', 0))
                row_cells[4].text = item.get('unit', '')
                row_cells[5].text = f"{item.get('our_rate', 0):,.2f}"
                row_cells[6].text = f"{item.get('total_amount', 0):,.2f}"
            
            # Total row
            total_row = table.add_row().cells
            total_row[0].text = ''
            total_row[1].text = 'TOTAL'
            total_row[1].merge(total_row[5])
            total_row[6].text = f"₹{data.get('total_our_value', 0):,.2f}"
            
            # Make total row bold
            for cell in total_row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph("Note: All prices are inclusive of GST @ 18%")
        doc.add_paragraph(f"Total Value: ₹{data.get('total_our_value', 0):,.2f}")
        
        # Save
        filename = f"boq_{data.get('tender_number', 'tender').replace('/', '_')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_compliance_statement(self, data: Dict[str, Any]) -> str:
        """
        Generate technical compliance statement
        
        Args:
            data: Technical requirements and compliance status
            
        Returns:
            Path to generated document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('Technical Compliance Statement', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Tender No: {data.get('tender_number', 'N/A')}")
        doc.add_paragraph()
        
        # Table
        requirements = data.get('technical_requirements', [])
        if requirements:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Clause'
            hdr_cells[1].text = 'Requirement'
            hdr_cells[2].text = 'Compliance'
            hdr_cells[3].text = 'Remarks'
            
            for req in requirements:
                row_cells = table.add_row().cells
                row_cells[0].text = req.get('clause_number', '')
                row_cells[1].text = req.get('requirement', '')
                row_cells[2].text = 'YES' if req.get('compliance', True) else 'NO'
                row_cells[3].text = req.get('remarks', 'Compliant')
        
        doc.add_paragraph()
        doc.add_paragraph("We hereby declare that we comply with all the above technical requirements.")
        
        # Save
        filename = f"compliance_{data.get('tender_number', 'tender').replace('/', '_')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath

# Global template engine instance
template_engine = DocumentTemplateEngine()
